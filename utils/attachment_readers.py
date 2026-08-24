from __future__ import annotations

import base64
import io
import json
from pathlib import Path
from typing import Any, Protocol
from uuid import uuid4

import pandas as pd
from defusedxml import ElementTree as SafeElementTree
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader
import xlrd

from utils.artifact_publication import is_published_artifact
from utils.llm_response import coerce_text_content
from utils.attachment_artifacts import (
    AttachmentError,
    AttachmentLimits,
    LocalAttachmentStore,
)
from utils.dataset_artifacts import (
    persist_dataset_artifact,
    portable_dataset_artifact,
)
from utils.user_storage import ThreadStorageScope


_MAX_STRUCTURE_DEPTH = 6
_MAX_STRUCTURE_ITEMS = 200


class VisionAnalyzer(Protocol):
    def analyze(self, *, mime: str, content: bytes, question: str) -> str: ...


class LangChainVisionAnalyzer:
    def __init__(self, llm: Any) -> None:
        self.llm = llm

    def analyze(self, *, mime: str, content: bytes, question: str) -> str:
        from langchain_core.messages import HumanMessage

        prompt = question.strip() or "Describe this attached image."
        encoded = base64.b64encode(content).decode("ascii")
        response = self.llm.invoke(
            [
                HumanMessage(
                    content=[
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime};base64,{encoded}",
                            },
                        },
                    ]
                )
            ]
        )
        return coerce_text_content(getattr(response, "content", response))


def _bounded_text(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[:limit], True


def _available_manifest(
    store: LocalAttachmentStore,
    thread_id: ThreadStorageScope | str,
    attachment_id: str,
) -> dict[str, Any]:
    manifest = store.require(thread_id, attachment_id)
    if manifest.get("status") not in {"available", "binding"}:
        raise AttachmentError(
            "ATTACHMENT_NOT_AVAILABLE",
            "attachment is not bound to an available conversation message",
        )
    return manifest


def _manifest_card(manifest: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(manifest.get("id") or ""),
        "filename": str(manifest.get("filename") or ""),
        "kind": str(manifest.get("kind") or ""),
        "format": str(manifest.get("format") or ""),
        "mime": str(manifest.get("mime") or ""),
        "byte_size": int(manifest.get("byte_size") or 0),
        "status": str(manifest.get("status") or ""),
    }


def _summarize_json(
    value: Any,
    *,
    depth: int = 0,
    budget: list[int] | None = None,
) -> Any:
    remaining = budget if budget is not None else [_MAX_STRUCTURE_ITEMS]
    if remaining[0] <= 0:
        return "…"
    remaining[0] -= 1
    if depth >= _MAX_STRUCTURE_DEPTH:
        return type(value).__name__
    if isinstance(value, dict):
        result: dict[str, Any] = {}
        for key, item in value.items():
            if remaining[0] <= 0:
                break
            result[str(key)[:200]] = _summarize_json(
                item,
                depth=depth + 1,
                budget=remaining,
            )
        return result
    if isinstance(value, list):
        return [
            _summarize_json(item, depth=depth + 1, budget=remaining)
            for item in value[:25]
            if remaining[0] > 0
        ]
    if isinstance(value, str):
        return value[:500]
    if isinstance(value, (int, float, bool)) or value is None:
        return value
    return str(value)[:500]


def _json_annotation_candidates(value: Any) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    if isinstance(value, dict):
        for key, item in list(value.items())[:_MAX_STRUCTURE_ITEMS]:
            candidate: dict[str, str] = {"name": str(key)[:200]}
            if isinstance(item, dict):
                for source, target in (
                    ("description", "description"),
                    ("label", "description"),
                    ("dataType", "dataType"),
                    ("type", "dataType"),
                ):
                    if source in item and target not in candidate:
                        candidate[target] = str(item[source])[:500]
            elif isinstance(item, str):
                candidate["description"] = item[:500]
            candidates.append(candidate)
    elif isinstance(value, list):
        for item in value[:_MAX_STRUCTURE_ITEMS]:
            if not isinstance(item, dict):
                continue
            name = item.get("name") or item.get("column") or item.get("id")
            if not name:
                continue
            candidate = {"name": str(name)[:200]}
            description = item.get("description") or item.get("label")
            if description:
                candidate["description"] = str(description)[:500]
            data_type = item.get("dataType") or item.get("type")
            if data_type:
                candidate["dataType"] = str(data_type)[:200]
            candidates.append(candidate)
    return candidates


def _xml_annotation_candidates(root: Any) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    for element in root.iter():
        if len(candidates) >= _MAX_STRUCTURE_ITEMS:
            break
        name = (
            element.attrib.get("name")
            or element.attrib.get("column")
            or element.attrib.get("id")
        )
        if not name:
            continue
        candidate = {"name": str(name)[:200]}
        text = " ".join(str(element.text or "").split())
        description = element.attrib.get("description") or text
        if description:
            candidate["description"] = str(description)[:500]
        data_type = element.attrib.get("dataType") or element.attrib.get("type")
        if data_type:
            candidate["dataType"] = str(data_type)[:200]
        candidates.append(candidate)
    return candidates


def _default_schema(dataframe: pd.DataFrame) -> dict[str, dict[str, str]]:
    return {
        str(column): {
            "description": "",
            "dataType": str(dataframe[column].dtype),
        }
        for column in dataframe.columns
    }


def _annotation_schema(
    dataframe: pd.DataFrame,
    annotations: list[dict[str, Any]],
) -> dict[str, dict[str, str]]:
    schema = _default_schema(dataframe)
    by_name: dict[str, dict[str, str]] = {}
    for annotation in annotations:
        for candidate in list(annotation.get("annotation_candidates") or []):
            if not isinstance(candidate, dict):
                continue
            name = str(candidate.get("name") or "")
            if name:
                by_name[name] = {
                    str(key): str(value)
                    for key, value in candidate.items()
                    if key in {"description", "dataType"} and value is not None
                }
    for column in schema:
        schema[column].update(by_name.get(column, {}))
    return schema


def _xml_dataframe(content: bytes, *, max_rows: int) -> pd.DataFrame:
    root = SafeElementTree.fromstring(content)
    rows: list[dict[str, Any]] = []
    for element in list(root):
        if len(rows) >= max_rows + 1:
            break
        row: dict[str, Any] = dict(element.attrib)
        children = list(element)
        if children:
            for child in children:
                row[str(child.tag)] = " ".join(str(child.text or "").split())
        else:
            text = " ".join(str(element.text or "").split())
            if text:
                row["value"] = text
        rows.append(row)
    if not rows:
        raise AttachmentError(
            "TABLE_PARSE_ERROR",
            "XML attachment does not contain repeated row elements",
        )
    return pd.DataFrame(rows)


def _validate_table_shape(
    *,
    rows: int,
    columns: int,
    limits: AttachmentLimits,
) -> None:
    if rows > limits.max_table_rows:
        raise AttachmentError(
            "TABLE_ROW_LIMIT",
            "table attachment exceeds the configured row limit",
        )
    if columns > limits.max_table_columns:
        raise AttachmentError(
            "TABLE_COLUMN_LIMIT",
            "table attachment exceeds the configured column limit",
        )
    if rows * columns > limits.max_table_cells:
        raise AttachmentError(
            "TABLE_CELL_LIMIT",
            "table attachment exceeds the configured decoded cell limit",
        )


def _json_records_dataframe(
    content: bytes,
    *,
    limits: AttachmentLimits,
) -> pd.DataFrame:
    text = content.decode("utf-8")
    decoder = json.JSONDecoder()
    position = 0
    while position < len(text) and text[position].isspace():
        position += 1
    if position >= len(text) or text[position] != "[":
        raise AttachmentError(
            "TABLE_PARSE_ERROR",
            "JSON table attachment must be an array of record objects",
        )
    position += 1
    records: list[dict[str, Any]] = []
    columns: set[str] = set()
    while position < len(text) and text[position].isspace():
        position += 1
    if position < len(text) and text[position] == "]":
        position += 1
        if text[position:].strip():
            raise AttachmentError(
                "TABLE_PARSE_ERROR",
                "JSON attachment contains content after the record array",
            )
        return pd.DataFrame()
    while True:
        if position < len(text) and text[position] == "]":
            raise AttachmentError(
                "TABLE_PARSE_ERROR",
                "JSON attachment contains a trailing comma",
            )
        try:
            value, position = decoder.raw_decode(text, position)
        except json.JSONDecodeError as exc:
            raise AttachmentError(
                "TABLE_PARSE_ERROR",
                "JSON attachment is not a valid record array",
            ) from exc
        if not isinstance(value, dict) or any(
            isinstance(item, (dict, list))
            for item in value.values()
        ):
            raise AttachmentError(
                "TABLE_PARSE_ERROR",
                "JSON table rows must be flat record objects",
            )
        records.append({str(key): item for key, item in value.items()})
        columns.update(str(key) for key in value)
        _validate_table_shape(
            rows=len(records),
            columns=len(columns),
            limits=limits,
        )
        while position < len(text) and text[position].isspace():
            position += 1
        if position < len(text) and text[position] == "]":
            position += 1
            break
        if position >= len(text) or text[position] != ",":
            raise AttachmentError(
                "TABLE_PARSE_ERROR",
                "JSON attachment is not a valid record array",
            )
        position += 1
        while position < len(text) and text[position].isspace():
            position += 1
    if text[position:].strip():
        raise AttachmentError(
            "TABLE_PARSE_ERROR",
            "JSON attachment contains content after the record array",
        )
    return pd.DataFrame(records)


def _read_dataframe(
    manifest: dict[str, Any],
    content: bytes,
    *,
    sheet_name: str | None,
    limits: AttachmentLimits,
) -> pd.DataFrame:
    file_format = manifest.get("format")
    try:
        if file_format == "csv":
            dataframe = pd.read_csv(
                io.BytesIO(content),
                nrows=limits.max_table_rows + 1,
            )
        elif file_format == "tsv":
            dataframe = pd.read_csv(
                io.BytesIO(content),
                sep="\t",
                nrows=limits.max_table_rows + 1,
            )
        elif file_format == "xlsx":
            workbook = load_workbook(
                io.BytesIO(content),
                read_only=True,
                data_only=True,
            )
            try:
                worksheet = (
                    workbook[sheet_name]
                    if sheet_name is not None
                    else workbook[workbook.sheetnames[0]]
                )
                _validate_table_shape(
                    rows=max(0, int(worksheet.max_row or 0) - 1),
                    columns=int(worksheet.max_column or 0),
                    limits=limits,
                )
            finally:
                workbook.close()
            dataframe = pd.read_excel(
                io.BytesIO(content),
                sheet_name=sheet_name if sheet_name is not None else 0,
                engine="openpyxl",
            )
        elif file_format == "xls":
            workbook = xlrd.open_workbook(
                file_contents=content,
                on_demand=True,
            )
            try:
                worksheet = (
                    workbook.sheet_by_name(sheet_name)
                    if sheet_name is not None
                    else workbook.sheet_by_index(0)
                )
                _validate_table_shape(
                    rows=max(0, int(worksheet.nrows) - 1),
                    columns=int(worksheet.ncols),
                    limits=limits,
                )
            finally:
                workbook.release_resources()
            dataframe = pd.read_excel(
                io.BytesIO(content),
                sheet_name=sheet_name if sheet_name is not None else 0,
                engine="xlrd",
            )
        elif file_format == "json":
            dataframe = _json_records_dataframe(
                content,
                limits=limits,
            )
        elif file_format == "xml":
            dataframe = _xml_dataframe(
                content,
                max_rows=limits.max_table_rows,
            )
        else:
            raise AttachmentError(
                "NOT_TABULAR",
                f"{manifest.get('filename')} is not a supported table attachment",
            )
    except AttachmentError:
        raise
    except Exception as exc:
        raise AttachmentError(
            "TABLE_PARSE_ERROR",
            f"Could not parse {manifest.get('filename')} as a table",
        ) from exc
    if not isinstance(dataframe, pd.DataFrame):
        raise AttachmentError(
            "TABLE_PARSE_ERROR",
            "Spreadsheet selection did not produce one table",
        )
    dataframe.columns = [str(column) for column in dataframe.columns]
    _validate_table_shape(
        rows=len(dataframe),
        columns=len(dataframe.columns),
        limits=limits,
    )
    if len(dataframe.columns) != len(set(dataframe.columns)):
        raise AttachmentError(
            "DUPLICATE_COLUMNS",
            "table attachment contains duplicate column names",
        )
    return dataframe


class AttachmentReaderService:
    def __init__(
        self,
        store: LocalAttachmentStore,
        runtime_root: str | Path,
        limits: AttachmentLimits | None = None,
        vision_analyzer: VisionAnalyzer | None = None,
        conversation_artifacts: dict[str, Any] | None = None,
    ) -> None:
        self.store = store
        self.runtime_root = Path(runtime_root).expanduser().resolve()
        self.limits = limits or store.limits
        self.vision_analyzer = vision_analyzer
        self.conversation_artifacts = dict(conversation_artifacts or {})

    def _conversation_record(
        self,
        attachment_id: str,
    ) -> tuple[str, dict[str, Any]] | None:
        file_record = dict(
            dict(self.conversation_artifacts.get("files") or {}).get(
                attachment_id
            )
            or {}
        )
        if file_record and is_published_artifact(file_record):
            return "file", file_record
        dataset = dict(
            dict(self.conversation_artifacts.get("datasets") or {}).get(
                attachment_id
            )
            or {}
        )
        if dataset.get("status") == "active":
            return "dataset", dataset
        return None

    def _available_manifest(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
    ) -> dict[str, Any]:
        try:
            return _available_manifest(self.store, thread_id, attachment_id)
        except AttachmentError as exc:
            if exc.code != "ATTACHMENT_NOT_FOUND":
                raise
        record = self._conversation_record(attachment_id)
        if record is None:
            raise AttachmentError(
                "ATTACHMENT_NOT_FOUND",
                "attachment was not found for this thread",
            )
        record_type, value = record
        mime = str(value.get("mime") or (
            "text/csv" if record_type == "dataset" else "application/octet-stream"
        ))
        kind = str(
            value.get("kind")
            or ("tabular" if record_type == "dataset" else "")
        )
        if mime.startswith("image/"):
            kind = "image"
        elif kind == "text":
            kind = "document"
        inferred_format = {
            "text/plain": "txt",
            "text/markdown": "md",
            "application/pdf": "pdf",
            "image/png": "png",
            "image/jpeg": "jpeg",
        }.get(mime, "")
        return {
            "id": attachment_id,
            "filename": str(
                value.get("filename")
                or value.get("summary")
                or value.get("name")
                or attachment_id
            ),
            "kind": kind,
            "format": (
                "csv"
                if record_type == "dataset"
                else str(value.get("format") or inferred_format)
            ),
            "mime": mime,
            "byte_size": int(value.get("byte_size") or 0),
            "status": "available",
        }

    def _read_bytes(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
    ) -> bytes:
        try:
            return self.store.read_bytes(thread_id, attachment_id)
        except AttachmentError as exc:
            if exc.code != "ATTACHMENT_NOT_FOUND":
                raise
        record = self._conversation_record(attachment_id)
        if record is None or record[0] != "file":
            raise AttachmentError(
                "ATTACHMENT_NOT_FOUND",
                "attachment content was not found for this thread",
            )
        content = record[1].get("content")
        if isinstance(content, str):
            return content.encode("utf-8")
        if isinstance(content, dict) and isinstance(
            content.get("data_base64"),
            str,
        ):
            return base64.b64decode(content["data_base64"])
        raise AttachmentError(
            "ATTACHMENT_NOT_FOUND",
            "attachment content was not found for this thread",
        )

    def inspect(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_ids: list[str],
    ) -> list[dict[str, Any]]:
        cards: list[dict[str, Any]] = []
        for attachment_id in attachment_ids:
            manifest = self._available_manifest(thread_id, attachment_id)
            if manifest.get("kind") == "tabular":
                cards.append(self.inspect_table_candidate(thread_id, attachment_id))
            else:
                cards.append(_manifest_card(manifest))
        return cards

    def inspect_table_candidate(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
    ) -> dict[str, Any]:
        manifest = self._available_manifest(thread_id, attachment_id)
        dataframe = _read_dataframe(
            manifest,
            self._read_bytes(thread_id, attachment_id),
            sheet_name=None,
            limits=self.limits,
        )
        sample_rows = json.loads(
            dataframe.head(5).to_json(orient="records", date_format="iso")
        )
        return {
            **_manifest_card(manifest),
            "columns": [str(column) for column in dataframe.columns[:100]],
            "row_count": int(len(dataframe)),
            "sample_rows": sample_rows,
        }

    def read_document(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
    ) -> dict[str, Any]:
        manifest = self._available_manifest(thread_id, attachment_id)
        content = self._read_bytes(thread_id, attachment_id)
        file_format = str(manifest.get("format") or "")
        page_count: int | None = None
        if file_format == "pdf":
            reader = PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            if page_count > self.limits.max_document_pages:
                raise AttachmentError(
                    "DOCUMENT_PAGE_LIMIT",
                    "PDF exceeds the configured document page limit",
                )
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif file_format == "docx":
            document = Document(io.BytesIO(content))
            text_parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    text_parts.append("\t".join(cell.text for cell in row.cells))
            text = "\n".join(text_parts)
        elif file_format in {"txt", "md"}:
            text = content.decode("utf-8")
        else:
            raise AttachmentError(
                "NOT_DOCUMENT",
                f"{manifest.get('filename')} is not a document attachment",
            )
        bounded, truncated = _bounded_text(
            text,
            self.limits.max_extracted_chars,
        )
        return {
            "id": attachment_id,
            "format": file_format,
            "text": bounded,
            "truncated": truncated,
            "page_count": page_count,
        }

    def parse_structured(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
    ) -> dict[str, Any]:
        manifest = self._available_manifest(thread_id, attachment_id)
        content = self._read_bytes(thread_id, attachment_id)
        file_format = str(manifest.get("format") or "")
        if file_format == "json":
            value = json.loads(content.decode("utf-8"))
            return {
                "id": attachment_id,
                "format": "json",
                "structure": _summarize_json(value),
                "annotation_candidates": _json_annotation_candidates(value),
            }
        if file_format == "xml":
            root = SafeElementTree.fromstring(content)
            return {
                "id": attachment_id,
                "format": "xml",
                "root": str(root.tag)[:200],
                "structure": {
                    "root": str(root.tag)[:200],
                    "child_tags": [
                        str(child.tag)[:200]
                        for child in list(root)[:_MAX_STRUCTURE_ITEMS]
                    ],
                },
                "annotation_candidates": _xml_annotation_candidates(root),
            }
        raise AttachmentError(
            "NOT_STRUCTURED",
            f"{manifest.get('filename')} is not a JSON or XML attachment",
        )

    def inspect_image(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
        *,
        question: str = "",
    ) -> dict[str, Any]:
        manifest = self._available_manifest(thread_id, attachment_id)
        if manifest.get("kind") != "image":
            raise AttachmentError(
                "NOT_IMAGE",
                f"{manifest.get('filename')} is not an image attachment",
            )
        content = self._read_bytes(thread_id, attachment_id)
        try:
            with Image.open(io.BytesIO(content)) as image:
                width, height = image.size
                image.verify()
        except Exception as exc:
            raise AttachmentError(
                "IMAGE_PARSE_ERROR",
                "image attachment could not be decoded",
            ) from exc
        if width * height > self.limits.max_image_pixels:
            raise AttachmentError(
                "IMAGE_PIXEL_LIMIT",
                "image exceeds the configured decoded pixel limit",
            )
        if self.vision_analyzer is None:
            raise AttachmentError(
                "VISION_UNAVAILABLE",
                "The configured model cannot inspect images",
            )
        description = self.vision_analyzer.analyze(
            mime=str(manifest.get("mime") or ""),
            content=content,
            question=question,
        )
        bounded, truncated = _bounded_text(
            str(description),
            self.limits.max_extracted_chars,
        )
        return {
            "id": attachment_id,
            "format": str(manifest.get("format") or ""),
            "width": width,
            "height": height,
            "description": bounded,
            "truncated": truncated,
        }

    def load_table(
        self,
        thread_id: ThreadStorageScope | str,
        attachment_id: str,
        *,
        sheet_name: str | None = None,
        annotation_ids: list[str] | None = None,
    ) -> dict[str, Any]:
        existing = self._conversation_record(attachment_id)
        if existing is not None and existing[0] == "dataset":
            return dict(existing[1])
        manifest = self._available_manifest(thread_id, attachment_id)
        dataframe = _read_dataframe(
            manifest,
            self._read_bytes(thread_id, attachment_id),
            sheet_name=sheet_name,
            limits=self.limits,
        )
        annotation_ids = list(annotation_ids or [])
        annotations = [
            self.parse_structured(thread_id, annotation_id)
            for annotation_id in annotation_ids
        ]
        source_attachment_ids = [attachment_id, *annotation_ids]
        source_manifests = [
            manifest,
            *[
                self._available_manifest(thread_id, annotation_id)
                for annotation_id in annotation_ids
            ],
        ]
        dataset_id = f"uploaded-{uuid4().hex[:12]}"
        dataset_root = (
            thread_id.datasets
            if isinstance(thread_id, ThreadStorageScope)
            else self.runtime_root
        )
        dataset_thread_id = (
            thread_id.thread_id
            if isinstance(thread_id, ThreadStorageScope)
            else thread_id
        )
        persisted = persist_dataset_artifact(
            runtime_root=dataset_root,
            thread_id=dataset_thread_id,
            dataset_id=dataset_id,
            kind="uploaded",
            dataframe=dataframe,
            schema=_annotation_schema(dataframe, annotations),
            provenance={
                "source": "message_attachment",
                "source_attachment_ids": source_attachment_ids,
                "source_filenames": [
                    str(source_manifest.get("filename") or "")
                    for source_manifest in source_manifests
                ],
            },
        )
        return portable_dataset_artifact(
            persisted,
            runtime_root=dataset_root,
        )
